import os
import re
import tempfile
from collections import OrderedDict
from io import BytesIO
from pathlib import Path

os.environ.setdefault(
    "MPLCONFIGDIR",
    str(Path(tempfile.gettempdir()) / "pilot_blind_test_matplotlib"),
)

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.io as pio
import streamlit as st
from openpyxl import load_workbook

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


st.set_page_config(page_title="飞行员双盲测试评估分析（空客）", layout="wide")


META_KEYWORDS = [
    "序号",
    "姓名",
    "日期",
    "所属单位",
    "所属 单位",
    "单位",
    "技术等级",
    "职务",
    "总飞行时间",
    "本机型经历时间",
    "检查员",
    "得分",
    "总得分",
]

SUBJECT_DEFS = [
    ("科目一", "大坡度盘旋"),
    ("科目二", "大侧风目视起落"),
    ("科目三", "选择的非精密进近"),
    ("科目四", "中断着陆+中断着陆后发动机失效"),
    ("科目五", "单发ILS无指引落地"),
    ("综合考评", "综合考评"),
]

SUBJECT_ORDER = [item[0] for item in SUBJECT_DEFS]
FLIGHT_SUBJECTS = SUBJECT_ORDER[:5]
SUBJECT_SORT_MAP = {subject_no: idx for idx, subject_no in enumerate(SUBJECT_ORDER, start=1)}
AIRBUS_STANDARD_DATA = [
    ("科目一", "高度偏差（偏差持续5秒，加倍）", "±20-40ft（含）", -1.0),
    ("科目一", "高度偏差（偏差持续5秒，加倍）", "±40-60ft（含）", -2.0),
    ("科目一", "高度偏差（偏差持续5秒，加倍）", "±60-80ft（含）", -3.0),
    ("科目一", "高度偏差（偏差持续5秒，加倍）", "±80-100ft（含）", -4.0),
    ("科目一", "速度偏差（偏差持续3秒，加倍）", "2-5KT（含）", -1.0),
    ("科目一", "速度偏差（偏差持续3秒，加倍）", "5-10KT（含）", -2.0),
    ("科目一", "速度偏差（偏差持续3秒，加倍）", "＞10KT", -3.0),
    ("科目一", "改出航向偏差", "2°-5°（含）", -1.0),
    ("科目一", "改出航向偏差", "5°-10°（含）", -2.0),
    ("科目一", "改出航向偏差", "＞±10°", -3.0),
    ("科目一", "坡度保持", "偏差持续3s以上", -1.0),
    ("科目一", "进入和改出滚转速率", "较明显停顿粗猛发力", -1.0),
    ("科目一", "进入和改出滚转速率", "明显停顿或粗猛发力", -2.0),
    ("科目二", "起飞抬头率", "过快或过慢", -1.0),
    ("科目二", "DON’T SINK", "DON’T SINK", -3.0),
    ("科目二", "离地后50ft飞机位置", "偏出跑道", -1.0),
    ("科目二", "一边航迹", "≥3°", -2.0),
    ("科目二", "一边/一转弯姿态", "≥25°", -3.0),
    ("科目二", "一转弯坡度", "＞30°", -3.0),
    ("科目二", "起始改平高度差", "＞100ft", -2.0),
    ("科目二", "三边高度改平后高度控制（偏差持续3秒，加倍）", "±40ft-±100ft（含）", -1.0),
    ("科目二", "三边高度改平后高度控制（偏差持续3秒，加倍）", "±100ft-±200ft（含）", -2.0),
    ("科目二", "三边高度改平后高度控制（偏差持续3秒，加倍）", "＞±200ft", -4.0),
    ("科目二", "三边航迹稳定后速度控制（偏差持续3秒计分）", "±5-±10KT（含）", -1.0),
    ("科目二", "三边航迹稳定后速度控制（偏差持续3秒计分）", "＞±10KT", -2.0),
    ("科目二", "三边宽度", "正切跑道头时＜2NM或＞3NM", -1.0),
    ("科目二", "四边航迹", "≥±5°", -1.0),
    ("科目二", "四转弯改出高度", "三红一白/三白一红", -1.0),
    ("科目二", "四转弯改出高度", "四白/四红", -2.0),
    ("科目二", "四转弯改出位置", "不在跑道内", -1.0),
    ("科目二", "下滑线（≥5秒，加倍）", "三红一白/三白一红", -1.0),
    ("科目二", "下滑线（≥5秒，加倍）", "四白/四红", -2.0),
    ("科目二", "五边速度控制（偏差持续3秒计分）", "±5-±10KT（含）", -1.0),
    ("科目二", "五边速度控制（偏差持续3秒计分）", "＞±10KT", -2.0),
    ("科目二", "五边位置", "500ft以下未飞入跑道延长线内", -1.0),
    ("科目二", "SINK RATE", "SINK RATE", -3.0),
    ("科目二", "入口高度", "相差每10ft", -1.0),
    ("科目二", "着陆", "曲线控制", -1.0),
    ("科目二", "着陆", "坡度≥5°", -1.0),
    ("科目二", "着陆", "偏流≥5°", -1.0),
    ("科目二", "着陆", "过早蹬舵（RA20ft以上）", -1.0),
    ("科目二", "着陆", "目测每偏差500ft", -1.0),
    ("科目二", "着陆", "超出接地区", -3.0),
    ("科目二", "着陆", "中心线不在两个主轮之间", -1.0),
    ("科目二", "着陆", "偏出大于半个跑道", -2.0),
    ("科目二", "着陆", "接地前抽杆", -1.0),
    ("科目二", "着陆", "拉平高后收光油门", -2.0),
    ("科目二", "着陆", "着陆弹跳", -2.0),
    ("科目二", "着陆", "主轮接地前拉出反推手柄", -3.0),
    ("科目二", "着陆", "明显过大的着陆载荷", -2.0),
    ("科目三", "滑跑方向", "不稳定", -1.0),
    ("科目三", "一边航迹", "≥3°", -1.0),
    ("科目三", "DON’T SINK", "DON’T SINK", -3.0),
    ("科目三", "扇区高度意识", "未考虑安全高度", -3.0),
    ("科目三", "位置管理意识", "未及时主动管理飞机位置", -1.0),
    ("科目三", "精确过台", "每偏差0.5NM", -1.0),
    ("科目三", "背台航迹控制", "超过3NM后，未在航迹上", -2.0),
    ("科目三", "程序转弯时机", "未按程序要求，提前/推迟转弯", -2.0),
    ("科目三", "四转弯方法", "过早或者过晚", -2.0),
    ("科目三", "稳定进近", "下降顶点未建立稳定进近", -3.0),
    ("科目三", "高距比", "是否参考原始数据", -3.0),
    ("科目三", "向台航迹控制", "≥半个点", -1.0),
    ("科目三", "下滑轨迹控制", "≥半个点", -1.0),
    ("科目三", "下滑线（≥5秒，加倍）", "三红一白/三白一红", -1.0),
    ("科目三", "下滑线（≥5秒，加倍）", "四白/四红", -2.0),
    ("科目三", "五边速度控制（偏差持续5秒计分）", "±5-±10KT（含）", -1.0),
    ("科目三", "五边速度控制（偏差持续5秒计分）", "＞±10KT", -2.0),
    ("科目三", "SINK RATE", "SINK RATE", -3.0),
    ("科目三", "仪表转目视飞行", "下降稳定性（PAPI变化）", -1.0),
    ("科目三", "仪表转目视飞行", "航迹稳定性（来回炒菜）", -1.0),
    ("科目三", "仪表转目视飞行", "坡度稳定性（单向变化超过5°）", -1.0),
    ("科目三", "横侧控制", "100ft以下超出跑道边线（含）", -1.0),
    ("科目三", "入口高度", "相差每10ft", -1.0),
    ("科目三", "入口位置", "不在跑道内", -5.0),
    ("科目四", "中断动作", "飞机姿态剧烈变化", -2.0),
    ("科目四", "中断动作", "接地载荷大", -2.0),
    ("科目四", "中断动作", "不合理使用推力", -1.0),
    ("科目四", "中断动作", "出现双输入", -2.0),
    ("科目四", "中断程序", "程序错误", -2.0),
    ("科目四", "程序", "单发初始姿态控制", -1.0),
    ("科目四", "程序", "单发推力控制", -1.0),
    ("科目四", "程序", "速度：＜Vref ＞Vref+20", -1.0),
    ("科目四", "离地位置控制", "50ft偏出跑道边线", -3.0),
    ("科目四", "航迹误差", "±5°-±10°（含）", -1.0),
    ("科目四", "航迹误差", "＞±10°", -2.0),
    ("科目四", "坡度", "≥5°", -1.0),
    ("科目四", "坡度", "≥10°", -2.0),
    ("科目四", "侧滑控制", "100ft侧滑未消除", -2.0),
    ("科目四", "侧滑控制", "未正确使用方向舵配平", -2.0),
    ("科目四", "通信", "没有宣布“PANPAN”", -1.0),
    ("科目五", "1000ft以下五边剖面控制", "1000ft未稳定进近", -1.0),
    ("科目五", "1000ft以下五边剖面控制", "LOC 每半个点", -1.0),
    ("科目五", "1000ft以下五边剖面控制", "G/S 每半个点", -1.0),
    ("科目五", "1000ft以下五边剖面控制", "100ft以下超出跑道边线（含）", -2.0),
    ("科目五", "速度控制", "Vapp：±5-±10KT（含）", -1.0),
    ("科目五", "速度控制", "Vapp：＞±10KT", -2.0),
    ("科目五", "SINK RATE", "SINK RATE", -3.0),
    ("科目五", "入口高度", "相差每10ft", -1.0),
    ("科目五", "方向舵配平", "未回中立位", -2.0),
    ("科目五", "着陆", "曲线控制", -1.0),
    ("科目五", "着陆", "坡度≥5°", -1.0),
    ("科目五", "着陆", "偏流≥5°", -1.0),
    ("科目五", "着陆", "过早蹬舵", -1.0),
    ("科目五", "着陆", "目测每偏差500ft", -1.0),
    ("科目五", "着陆", "超出接地区", -3.0),
    ("科目五", "着陆", "中心线不在两个主轮之间", -1.0),
    ("科目五", "着陆", "偏出大于半个跑道", -2.0),
    ("科目五", "着陆", "接地前抽杆", -1.0),
    ("科目五", "着陆", "着陆弹跳", -2.0),
    ("科目五", "着陆", "明显过大的着陆载荷", -2.0),
    ("综合考评", "综合考评", "冲偏出跑道", -20.0),
    ("综合考评", "综合考评", "飞机结构受损", -20.0),
    ("综合考评", "综合考评", "跑道外接地", -20.0),
    ("综合考评", "综合考评", "盲目蛮干", -20.0),
    ("综合考评", "综合考评", "PULL UP", -10.0),
    ("综合考评", "综合考评", "技术性复飞", -10.0),
    ("综合考评", "综合考评", "决策失误", -10.0),
    ("综合考评", "综合考评", "各类音响警戒/警告", -3.0),
    ("综合考评", "综合考评", "算错Vapp", -2.0),
]
AIRBUS_RISK_ITEM_MAP = {
    ("科目一", "高度偏差（偏差持续5秒，加倍）"): "飞行高度突破MAC03",
    ("科目一", "速度偏差（偏差持续3秒，加倍）"): "飞机操纵RE16",
    ("科目一", "改出航向偏差"): "横向偏离MAC02",
    ("科目一", "坡度保持"): "飞机操纵RE16",
    ("科目一", "进入和改出滚转速率"): "飞机操纵RE16",
    ("科目二", "起飞抬头率"): "起飞抬头率RE08",
    ("科目二", "DON’T SINK"): "地形警告FIT08",
    ("科目二", "离地后50ft飞机位置"): "横向偏离MAC02",
    ("科目二", "一边航迹"): "横向偏离MAC02",
    ("科目二", "一边/一转弯姿态"): "不正常姿态LOC14",
    ("科目二", "一转弯坡度"): "不正常姿态LOC14",
    ("科目二", "起始改平高度差"): "飞行高度突破MAC03",
    ("科目二", "三边高度改平后高度控制（偏差持续3秒，加倍）"): "飞行高度突破MAC03",
    ("科目二", "三边航迹稳定后速度控制（偏差持续3秒计分）"): "飞机操纵RE16",
    ("科目二", "三边宽度"): "横向偏离MAC02",
    ("科目二", "四边航迹"): "横向偏离MAC02",
    ("科目二", "四转弯改出高度"): "偏离下滑道下方CFIT04",
    ("科目二", "四转弯改出位置"): "横向偏离RE20",
    ("科目二", "下滑线（≥5秒，加倍）"): "偏离下滑道下方CFIT04",
    ("科目二", "五边速度控制（偏差持续3秒计分）"): "着陆时能量过高RE32",
    ("科目二", "五边位置"): "横向偏离MAC02",
    ("科目二", "SINK RATE"): "进近阶段低能量状态/不稳定进近CFIT11",
    ("科目二", "入口高度"): "跑道入口处能量过高RE27",
    ("科目二", "着陆"): "飞机操纵RE16",
    ("科目三", "滑跑方向"): "横向偏离RE20",
    ("科目三", "一边航迹"): "横向偏离MAC02",
    ("科目三", "DON’T SINK"): "地形警告FIT08",
    ("科目三", "扇区高度意识"): "横侧距离障碍物过近CFIT13",
    ("科目三", "位置管理意识"): "横向偏离MAC02",
    ("科目三", "精确过台"): "横向偏离MAC02",
    ("科目三", "背台航迹控制"): "横侧距离障碍物过近CFIT13",
    ("科目三", "程序转弯时机"): "横侧距离障碍物过近CFIT13",
    ("科目三", "四转弯方法"): "横侧距离障碍物过近CFIT13",
    ("科目三", "稳定进近"): "不稳定进近RE26",
    ("科目三", "高距比"): "错误下降点CFIT07",
    ("科目三", "向台航迹控制"): "横侧距离障碍物过近CFIT13",
    ("科目三", "下滑轨迹控制"): "偏离下滑道下方CFIT04",
    ("科目三", "下滑线（≥5秒，加倍）"): "偏离下滑道下方CFIT04",
    ("科目三", "五边速度控制（偏差持续5秒计分）"): "着陆时能量过高RE32",
    ("科目三", "SINK RATE"): "进近阶段低能量状态/不稳定进近CFIT11",
    ("科目三", "仪表转目视飞行"): "不正常姿态LOC14",
    ("科目三", "横侧控制"): "横向偏离RE20",
    ("科目三", "入口高度"): "跑道入口处能量过高RE27",
    ("科目三", "入口位置"): "横向偏离RE20",
    ("科目四", "中断动作"): "复飞RE31",
    ("科目四", "中断程序"): "复飞RE31",
    ("科目四", "程序"): "发动机故障LOC23",
    ("科目四", "离地位置控制"): "横向偏离RE20",
    ("科目四", "航迹误差"): "横向偏离MAC02",
    ("科目四", "坡度"): "飞机操纵RE16",
    ("科目四", "侧滑控制"): "横向偏离RE20",
    ("科目四", "通信"): "发动机故障LOC23",
    ("科目五", "1000ft以下五边剖面控制"): "不稳定进近RE26",
    ("科目五", "速度控制"): "着陆时能量过高RE32",
    ("科目五", "SINK RATE"): "进近阶段低能量状态/不稳定进近CFIT11",
    ("科目五", "入口高度"): "跑道入口处能量过高RE27",
    ("科目五", "方向舵配平"): "横向偏离RE20",
    ("科目五", "着陆"): "飞机操纵RE16",
}
AIRBUS_RISK_STANDARD_MAP = {
    ("科目四", "中断动作", "出现双输入"): "不正常姿态LOC14",
}
SUBJECT_START_ITEMS = OrderedDict(
    [
        ("高度偏差", ("科目一", "大坡度盘旋")),
        ("起飞抬头率", ("科目二", "大侧风目视起落")),
        ("滑跑方向", ("科目三", "非精密进近+中断着陆")),
        ("中断动作", ("科目四", "中断着陆后发动机失效")),
        ("程序", ("科目四", "中断着陆后发动机失效")),
        ("1000ft以下五边剖面控制", ("科目五", "单发ILS无指引落地")),
        ("综合考评", ("综合考评", "综合考评")),
    ]
)
DEFAULT_COLOR_SEQUENCE = [
    "#2F80ED",
    "#27AE60",
    "#F2994A",
    "#9B51E0",
    "#EB5757",
    "#00A3A3",
    "#828282",
]
CAPTAIN_COLOR = "#27AE60"
FIRST_OFFICER_COLOR = "#2F80ED"
OVERALL_SCORE_COLOR = "#5DADE2"
AVERAGE_LINE_COLOR = "#F2994A"

# 科目颜色映射（蓝绿橙紫红）
SUBJECT_COLORS = {
    "科目一": "#2F80ED",  # 蓝色
    "科目二": "#27AE60",  # 绿色
    "科目三": "#F2994A",  # 橙色
    "科目四": "#9B51E0",  # 紫色
    "科目五": "#EB5757",  # 红色
}


def compact_text(value):
    if value is None or pd.isna(value):
        return ""
    text = str(value).replace("\n", " ").replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace(" （", "（").replace(" ）", "）")
    return text


def key_text(value):
    return re.sub(r"\s+", "", compact_text(value))


def risk_name_for_standard(subject_no, scoring_item, standard, fallback=""):
    subject_no = compact_text(subject_no)
    scoring_item = compact_text(scoring_item)
    standard = compact_text(standard)
    return (
        AIRBUS_RISK_STANDARD_MAP.get((subject_no, scoring_item, standard))
        or AIRBUS_RISK_ITEM_MAP.get((subject_no, scoring_item))
        or fallback
        or ""
    )


def to_number(value):
    if value is None or value == "":
        return np.nan
    if isinstance(value, str) and value.startswith("="):
        return np.nan
    return pd.to_numeric(value, errors="coerce")


def make_unique_name(name, counts):
    base = name or "unnamed"
    counts[base] = counts.get(base, 0) + 1
    return base if counts[base] == 1 else f"{base}_{counts[base]}"


def subject_from_item(item_name, current_subject):
    item_key = key_text(item_name)
    for marker, subject in SUBJECT_START_ITEMS.items():
        if item_key == key_text(marker):
            return subject
    return current_subject or SUBJECT_DEFS[0]


def subject_sort_value(subject_no):
    return SUBJECT_SORT_MAP.get(subject_no, len(SUBJECT_SORT_MAP) + 1)


def normalize_role(value):
    text = compact_text(value)
    if "副" in text:
        return "副驾驶"
    return "机长"


def useful_company_from_filename(file_name):
    stem = Path(file_name).stem
    text = re.sub(r"[（）()【】\[\]]", "_", stem)
    text = re.sub(r"\d{4}[.-]?\d{0,2}[.-]?\d{0,2}", "_", text)
    text = re.sub(r"\b(?:A|B|C|ARJ|C919|C909)\s*\d{2,4}[A-Z]?\b", "_", text, flags=re.I)
    text = re.sub(r"(?:B737|A320|C909|C919|ARJ21|ARJ|COMAC)", "_", text, flags=re.I)

    noise = [
        "华东飞行员双盲测试数据采集表",
        "飞行员双盲测试数据采集表",
        "双盲测试数据采集表",
        "华东局",
        "飞行技能评估标准",
        "双盲测试",
        "数据采集表",
        "评分表",
        "测试",
        "原始数据",
        "R2",
    ]
    for word in noise:
        text = text.replace(word, "_")

    candidates = [c.strip("_- 　") for c in re.split(r"[_\s\-]+", text) if c.strip("_- 　")]
    cleaned = []
    for item in candidates:
        item = re.sub(r"(?:机型|航班|版本|模板|汇总)$", "", item)
        if len(item) >= 2 and not re.fullmatch(r"[A-Za-z0-9.]+", item):
            cleaned.append(item)

    if cleaned:
        return cleaned[-1]
    return stem


def first_existing_column(df, names):
    lookup = {key_text(col): col for col in df.columns}
    for name in names:
        key = key_text(name)
        if key in lookup:
            return lookup[key]
    for col in df.columns:
        col_key = key_text(col)
        if any(key_text(name) in col_key for name in names):
            return col
    return None


def fill_id(cell):
    color = cell.fill.fgColor
    return color.rgb or color.indexed or color.theme or ""


def is_deduction_column(ws, col_idx):
    row2 = compact_text(ws.cell(2, col_idx).value)
    row3 = to_number(ws.cell(3, col_idx).value)
    if not np.isnan(row3) and row3 < 0:
        return True
    if col_idx >= 10 and row2 and key_text(ws.cell(1, col_idx).value) not in {"总得分", "得分"}:
        return True
    return False


def build_column_metadata(ws):
    current_item = ""
    current_subject = SUBJECT_DEFS[0]
    standard_idx = 0
    counts = {}
    columns = []
    col_meta = OrderedDict()
    deduction_cols = []

    max_col = 1
    for col_idx in range(1, ws.max_column + 1):
        if any(ws.cell(row_idx, col_idx).value not in (None, "") for row_idx in (1, 2, 3)):
            max_col = col_idx

    for col_idx in range(1, max_col + 1):
        top = compact_text(ws.cell(1, col_idx).value)
        bottom = compact_text(ws.cell(2, col_idx).value)
        score = to_number(ws.cell(3, col_idx).value)

        if top:
            current_item = top

        if is_deduction_column(ws, col_idx):
            if standard_idx < len(AIRBUS_STANDARD_DATA):
                subject_no, scoring_item, deduction_item, score = AIRBUS_STANDARD_DATA[standard_idx]
                subject_name = dict(SUBJECT_DEFS).get(subject_no, subject_no)
                column_order = standard_idx + 1
                standard_idx += 1
            else:
                scoring_item = current_item or top or "未命名评分项目"
                current_subject = subject_from_item(scoring_item, current_subject)
                subject_no, subject_name = current_subject
                deduction_item = bottom or scoring_item
                column_order = len(deduction_cols) + 1
            col_name = f"{subject_no}_{scoring_item}_{deduction_item}"
            col_name = make_unique_name(col_name, counts)
            deduction_cols.append(col_name)
            col_meta[col_name] = {
                "列号": col_idx,
                "列顺序": column_order,
                "科目排序": subject_sort_value(subject_no),
                "科目编号": subject_no,
                "科目名称": subject_name,
                "评分项目": scoring_item,
                "扣分标准": deduction_item,
                "标准分值": float(score) if not np.isnan(score) else np.nan,
                "扣分项": col_name,
            }
        else:
            raw_name = top or bottom or f"unnamed_{col_idx}"
            normalized = key_text(raw_name)
            if normalized == "所属单位":
                raw_name = "所属单位"
            col_name = make_unique_name(compact_text(raw_name), counts)
            col_meta[col_name] = {
                "列号": col_idx,
                "科目编号": "",
                "科目名称": "",
                "评分项目": col_name,
                "扣分标准": "",
                "标准分值": np.nan,
                "扣分项": col_name,
            }

        columns.append(col_name)

    return columns, col_meta, deduction_cols


@st.cache_data(show_spinner=False)
def load_template(file_bytes, file_name):
    try:
        wb = load_workbook(BytesIO(file_bytes), data_only=True)
        ws = wb[wb.sheetnames[0]]
        columns, col_meta, deduction_cols = build_column_metadata(ws)

        metadata_cols = [
            col
            for col in columns
            if any(key_text(word) == key_text(col) or key_text(word) in key_text(col) for word in META_KEYWORDS)
        ]
        fill_down_cols = [
            col
            for col in metadata_cols
            if not any(word in key_text(col) for word in ["检查员", "得分", "总得分"])
        ]

        records = []
        last_values = {col: "" for col in fill_down_cols}
        source_unit = useful_company_from_filename(file_name)

        for row_idx in range(4, ws.max_row + 1):
            row = {}
            has_any = False
            for col_idx, col_name in enumerate(columns, start=1):
                value = ws.cell(row_idx, col_idx).value
                if value not in (None, ""):
                    has_any = True
                row[col_name] = value

            if not has_any:
                continue

            for col in fill_down_cols:
                if row.get(col) not in (None, ""):
                    last_values[col] = row[col]
                else:
                    row[col] = last_values.get(col, "")

            row["来源文件"] = file_name
            row["文件单位"] = source_unit
            records.append(row)

        if not records:
            return pd.DataFrame(), col_meta, deduction_cols

        df = pd.DataFrame(records)

        name_col = first_existing_column(df, ["姓名"])
        if name_col:
            df = df[df[name_col].notna() & (df[name_col].astype(str).str.strip() != "")]

        unit_col = first_existing_column(df, ["所属单位", "单位"])
        if unit_col and unit_col in df.columns:
            internal_unit = df[unit_col].dropna().astype(str).map(compact_text)
            internal_unit = internal_unit[internal_unit != ""]
            if source_unit == Path(file_name).stem and not internal_unit.empty:
                source_unit = internal_unit.mode().iloc[0]
        df["所属单位"] = source_unit

        seq_col = first_existing_column(df, ["序号"])
        if seq_col:
            seq = df[seq_col].astype(str).map(compact_text)
        else:
            seq = pd.Series(range(1, len(df) + 1), index=df.index).astype(str)
        name = df[name_col].astype(str).map(compact_text) if name_col else seq
        df["人员ID"] = df["来源文件"].astype(str) + "::" + seq + "::" + name

        return df.reset_index(drop=True), col_meta, deduction_cols
    except Exception as exc:
        st.warning(f"{file_name} 解析失败：{exc}")
        return pd.DataFrame(), {}, []


def prepare_analysis(raw_data, col_meta, deduction_cols):
    df = raw_data.copy()
    deduction_cols = [col for col in deduction_cols if col in df.columns]

    if not deduction_cols:
        df["扣分总和"] = 0.0
        df["失分"] = 0.0
        df["最终得分"] = 100.0
        df["扣分项数量"] = 0
        return df, df.copy(), pd.DataFrame(), pd.DataFrame(), {}

    numeric = df[deduction_cols].apply(pd.to_numeric, errors="coerce").fillna(0)
    negative = numeric.where(numeric < 0, 0)
    df["扣分总和"] = negative.sum(axis=1)
    df["失分"] = -df["扣分总和"]
    df["最终得分"] = 100 + df["扣分总和"]
    df["扣分项数量"] = (negative < 0).sum(axis=1)

    name_col = first_existing_column(df, ["姓名"])
    tech_col = first_existing_column(df, ["技术等级", "职务"])
    inspector_col = first_existing_column(df, ["检查员"])
    score_col = first_existing_column(df, ["得分", "总得分"])

    df["姓名"] = df[name_col].map(compact_text) if name_col else ""
    df["技术等级"] = df[tech_col].map(normalize_role) if tech_col else "机长"
    df["操纵者"] = df["技术等级"]
    df["检查员"] = df[inspector_col].map(compact_text) if inspector_col else ""
    df["模板得分"] = pd.to_numeric(df[score_col], errors="coerce") if score_col else np.nan
    df["是否合计行"] = df["检查员"].str.contains("合计", na=False)
    df["是否平均行"] = df["检查员"].str.contains("平均|合计", na=False)

    checker_df = df[~df["是否平均行"]].copy()
    if checker_df.empty:
        checker_df = df.copy()

    agg_map = {
        "姓名": "first",
        "所属单位": "first",
        "技术等级": "first",
        "操纵者": "first",
        "最终得分": "mean",
        "扣分总和": "mean",
        "失分": "mean",
        "扣分项数量": "mean",
        "检查员": "count",
    }
    pilot_df = (
        checker_df.groupby("人员ID", dropna=False)
        .agg(agg_map)
        .rename(columns={"检查员": "评分人数"})
        .reset_index()
    )
    total_score_df = (
        df[df["是否合计行"]]
        .groupby("人员ID", dropna=False)["模板得分"]
        .first()
        .reset_index()
        .rename(columns={"模板得分": "合计得分"})
    )
    if not total_score_df.empty:
        pilot_df = pilot_df.merge(total_score_df, on="人员ID", how="left")
        pilot_df["最终得分"] = pilot_df["合计得分"].combine_first(pilot_df["最终得分"])
        pilot_df["扣分总和"] = pilot_df["最终得分"] - 100
        pilot_df["失分"] = 100 - pilot_df["最终得分"]
        pilot_df = pilot_df.drop(columns=["合计得分"])
    pilot_df["最终得分"] = pilot_df["最终得分"].round(2)
    pilot_df["平均失分"] = pilot_df["失分"].round(2)
    pilot_df["平均扣分项数量"] = pilot_df["扣分项数量"].round(2)

    subject_to_cols = OrderedDict()
    ordered_deduction_cols = sorted(
        deduction_cols,
        key=lambda col: (
            col_meta.get(col, {}).get("科目排序", subject_sort_value(col_meta.get(col, {}).get("科目编号", ""))),
            col_meta.get(col, {}).get("列顺序", 9999),
        ),
    )
    for col in ordered_deduction_cols:
        subject_no = col_meta.get(col, {}).get("科目编号", "未识别科目")
        subject_to_cols.setdefault(subject_no, []).append(col)

    for subject_no, cols in subject_to_cols.items():
        checker_df[f"{subject_no}_失分"] = -negative.loc[checker_df.index, cols].sum(axis=1)

    subject_loss_cols = [f"{subject_no}_失分" for subject_no in subject_to_cols]
    subject_avg = (
        checker_df.groupby("人员ID", dropna=False)[subject_loss_cols]
        .mean()
        .reset_index()
        if subject_loss_cols
        else pd.DataFrame()
    )
    pilot_subject = pilot_df[["人员ID", "姓名", "所属单位", "技术等级", "操纵者"]].merge(
        subject_avg, on="人员ID", how="left"
    )

    subject_melt = pd.DataFrame()
    if subject_loss_cols and not pilot_subject.empty:
        subject_melt = pilot_subject.melt(
            id_vars=["人员ID", "姓名", "所属单位", "技术等级", "操纵者"],
            value_vars=subject_loss_cols,
            var_name="科目编号",
            value_name="平均失分",
        )
        subject_melt["科目编号"] = subject_melt["科目编号"].str.replace("_失分", "", regex=False)
        subject_name_map = {no: name for no, name in SUBJECT_DEFS}
        subject_melt["科目名称"] = subject_melt["科目编号"].map(subject_name_map).fillna(subject_melt["科目编号"])
        subject_melt["科目显示"] = np.where(
            subject_melt["科目编号"].isin(FLIGHT_SUBJECTS),
            subject_melt["科目编号"] + "_" + subject_melt["科目名称"],
            subject_melt["科目编号"],
        )
        subject_melt["科目得分"] = 100 - subject_melt["平均失分"].fillna(0)

    records = []
    for col in ordered_deduction_cols:
        vals = pd.to_numeric(checker_df[col], errors="coerce").fillna(0)
        negatives = vals[vals < 0]
        meta = col_meta.get(col, {})
        for idx, val in negatives.items():
            row = checker_df.loc[idx]
            records.append(
                {
                    "记录ID": idx,
                    "人员ID": row.get("人员ID", ""),
                    "姓名": row.get("姓名", ""),
                    "所属单位": row.get("所属单位", ""),
                    "技术等级": row.get("技术等级", ""),
                    "操纵者": row.get("操纵者", ""),
                    "检查员": row.get("检查员", ""),
                    "科目编号": meta.get("科目编号", "未识别科目"),
                    "科目排序": meta.get("科目排序", subject_sort_value(meta.get("科目编号", ""))),
                    "科目名称": meta.get("科目名称", ""),
                    "评分项目": meta.get("评分项目", ""),
                    "扣分标准": meta.get("扣分标准", ""),
                    "扣分项": meta.get("扣分项", col),
                    "模板列": col,
                    "列顺序": meta.get("列顺序", 9999),
                    "标准分值": meta.get("标准分值", np.nan),
                    "扣分值": float(val),
                    "失分": float(abs(val)),
                }
            )
    deductions = pd.DataFrame(records)
    if not deductions.empty:
        deductions["科目显示"] = np.where(
            deductions["科目编号"].isin(FLIGHT_SUBJECTS),
            deductions["科目编号"] + "_" + deductions["科目名称"],
            deductions["科目编号"],
        )

    return df, pilot_df, deductions, subject_melt, subject_to_cols


def all_deduction_items(col_meta, subject_no=None):
    rows = []
    for col, meta in col_meta.items():
        item_subject = meta.get("科目编号", "")
        if not item_subject:
            continue
        if subject_no is not None and item_subject != subject_no:
            continue
        rows.append(
            {
                "扣分项": meta.get("扣分项", ""),
                "科目编号": item_subject,
                "科目排序": meta.get("科目排序", subject_sort_value(item_subject)),
                "列顺序": meta.get("列顺序", 9999),
                "评分项目": meta.get("评分项目", ""),
                "扣分标准": meta.get("扣分标准", ""),
                "标准分值": meta.get("标准分值", np.nan),
                "模板列": col,
            }
        )
    if not rows:
        return pd.DataFrame(columns=["扣分项", "科目编号", "科目排序", "列顺序", "评分项目", "扣分标准", "标准分值", "模板列"])
    return (
        pd.DataFrame(rows)
        .drop_duplicates(subset=["模板列"])
        .sort_values(["科目排序", "列顺序"])
        .reset_index(drop=True)
    )


def aggregate_loss_by_item(deductions, group_cols=None, sort_by_loss=True, all_items=None):
    group_cols = group_cols or ["扣分项"]
    if deductions.empty:
        out = pd.DataFrame(columns=group_cols + ["总失分", "扣分次数", "科目排序", "列顺序"])
    else:
        out = (
            deductions.groupby(group_cols, dropna=False)
            .agg(
                总失分=("失分", "sum"),
                扣分次数=("扣分项", "size"),
                科目排序=("科目排序", "min"),
                列顺序=("列顺序", "min"),
            )
            .reset_index()
        )
    if all_items is not None and group_cols == ["扣分项"] and not all_items.empty:
        base = all_items[["扣分项", "科目排序", "列顺序"]].drop_duplicates("扣分项")
        out = base.merge(
            out.drop(columns=[col for col in ["科目排序", "列顺序"] if col in out.columns]),
            on="扣分项",
            how="left",
        )
        out["总失分"] = out["总失分"].fillna(0)
        out["扣分次数"] = out["扣分次数"].fillna(0).astype(int)
    if "总失分" in out.columns:
        out["总扣分值"] = negative_deduction_values(out["总失分"])
    if sort_by_loss:
        return out.sort_values(["总失分", "扣分次数"], ascending=[False, False])
    return out.sort_values(["科目排序", "列顺序"])


def identify_weak_areas(deductions, group_cols=None, denominator=1):
    if deductions.empty:
        return pd.DataFrame()
    group_cols = group_cols or ["扣分项"]
    out = (
        deductions.groupby(group_cols, dropna=False)
        .agg(
            扣分次数=("扣分项", "size"),
            总失分=("失分", "sum"),
            平均单次失分=("失分", "mean"),
            科目排序=("科目排序", "min"),
            列顺序=("列顺序", "min"),
        )
        .reset_index()
    )
    out["出现率"] = (out["扣分次数"] / max(denominator, 1) * 100).round(1).astype(str) + "%"
    out["总扣分值"] = negative_deduction_values(out["总失分"])
    return out.sort_values(["总失分", "扣分次数"], ascending=[False, False])


def company_stats(pilot_df):
    if pilot_df.empty:
        return pd.DataFrame()
    return (
        pilot_df.groupby("所属单位")["最终得分"]
        .agg(人数="count", 平均分="mean", 最高分="max", 最低分="min", 标准差="std")
        .round(2)
        .reset_index()
        .rename(columns={"所属单位": "单位名称"})
    )


def company_test_counts(pilot_df):
    """按数据文件中的受测人员口径统计各航司人数。"""
    if pilot_df.empty:
        return pd.DataFrame(columns=["所属单位", "测试人数"])
    return (
        pilot_df.groupby("所属单位", dropna=False)["人员ID"]
        .nunique()
        .reset_index(name="测试人数")
    )


def display_loss_table(df):
    if df.empty:
        return df
    return df.drop(columns=["总失分"], errors="ignore")


def negative_deduction_values(values):
    values = pd.to_numeric(values, errors="coerce").fillna(0)
    return values.where(values == 0, -values)


def figure_height(rows, minimum=420, per_row=30, maximum=900):
    return int(min(max(minimum, rows * per_row + 180), maximum))


def fig_score_distribution(df, title="测试人员平均得分分布"):
    """显示每个飞行员的实际得分分布，按分数从低到高排列"""
    if df.empty:
        return None
    
    # 准备数据：按得分排序
    plot_df = df.copy()
    plot_df = plot_df.sort_values("最终得分", ascending=True).reset_index(drop=True)
    plot_df["人员序号"] = plot_df.index + 1
    
    # 创建条形图显示每个飞行员的得分
    fig = px.bar(
        plot_df,
        x="人员序号",
        y="最终得分",
        title=title,
        labels={"人员序号": "飞行员（按得分从低到高排列）", "最终得分": "平均得分"},
        text=plot_df["最终得分"].map(lambda x: f"{x:.2f}"),
    )
    fig.update_traces(marker_color=OVERALL_SCORE_COLOR)
    
    # 添加参考线：平均分
    avg_score = df["最终得分"].mean()
    fig.add_hline(
        y=avg_score, 
        line_dash="dash", 
        line_color=AVERAGE_LINE_COLOR,
    )
    fig.add_annotation(
        x=1,
        xref="paper",
        y=avg_score,
        yref="y",
        text=f"平均分: {avg_score:.2f}",
        showarrow=False,
        xanchor="right",
        yanchor="bottom",
        yshift=8,
        bgcolor="rgba(255,255,255,0.92)",
        bordercolor=AVERAGE_LINE_COLOR,
        borderwidth=1,
        font=dict(color="#333333", size=12),
    )
    
    fig.update_traces(
        textposition="outside",
        cliponaxis=False,
        hovertemplate="<b>%{customdata}</b><br>得分: %{y:.2f}<extra></extra>",
        customdata=plot_df["姓名"] if "姓名" in plot_df.columns else None
    )
    
    # 设置y轴范围
    y_min = max(0, plot_df["最终得分"].min() - 5)
    y_max = min(105, max(plot_df["最终得分"].max(), avg_score) + 8)
    
    fig.update_layout(
        height=430,
        xaxis_title="飞行员（按得分从低到高排列）",
        yaxis_title="平均得分",
        yaxis=dict(
            range=[y_min, y_max],
            dtick=5
        ),
        margin=dict(l=40, r=40, t=50, b=60)
    )
    
    return fig


def fig_score_distribution_by_role(df, role_type="机长", color_scale=None):
    """按角色（机长/副驾驶）显示分数分布"""
    if df.empty:
        return None
    
    # 筛选数据
    role_df = df[df["操纵者"] == role_type].copy()
    if role_df.empty:
        return None
    
    # 按得分排序
    role_df = role_df.sort_values("最终得分", ascending=True).reset_index(drop=True)
    role_df["人员序号"] = role_df.index + 1
    
    role_color = CAPTAIN_COLOR if role_type == "机长" else FIRST_OFFICER_COLOR
    
    # 创建条形图
    fig = px.bar(
        role_df,
        x="人员序号",
        y="最终得分",
        title=f"{role_type}分数分布",
        labels={"人员序号": f"{role_type}（按得分从低到高排列）", "最终得分": "平均得分"},
        text=role_df["最终得分"].map(lambda x: f"{x:.2f}"),
    )
    fig.update_traces(marker_color=role_color)
    
    # 添加平均分参考线
    avg_score = role_df["最终得分"].mean()
    fig.add_hline(
        y=avg_score,
        line_dash="dash",
        line_color=AVERAGE_LINE_COLOR,
    )
    fig.add_annotation(
        x=1,
        xref="paper",
        y=avg_score,
        yref="y",
        text=f"平均分: {avg_score:.2f}",
        showarrow=False,
        xanchor="right",
        yanchor="bottom",
        yshift=8,
        bgcolor="rgba(255,255,255,0.92)",
        bordercolor=AVERAGE_LINE_COLOR,
        borderwidth=1,
        font=dict(color="#333333", size=12),
    )
    
    fig.update_traces(
        textposition="outside",
        cliponaxis=False,
        hovertemplate="<b>%{customdata}</b><br>得分: %{y:.2f}<extra></extra>",
        customdata=role_df["姓名"] if "姓名" in role_df.columns else None
    )
    
    # 设置y轴范围
    y_min = max(0, role_df["最终得分"].min() - 5)
    y_max = min(105, max(role_df["最终得分"].max(), avg_score) + 8)
    
    fig.update_layout(
        height=400,
        xaxis_title=f"{role_type}（按得分从低到高排列）",
        yaxis_title="平均得分",
        yaxis=dict(
            range=[y_min, y_max],
            dtick=5
        ),
        margin=dict(l=40, r=40, t=60, b=50)
    )
    
    return fig


def fig_participants_by_company(pilot_df):
    counts = (
        pilot_df.groupby(["所属单位", "操纵者"], dropna=False)
        .size()
        .reset_index(name="人数")
        .sort_values(["所属单位", "操纵者"])
    )
    # 绘图顺序用于堆叠：副驾驶先画在下方，机长后画在上方；图例再单独排序。
    stack_order = ["副驾驶", "机长"]
    legend_order = ["机长", "副驾驶"]
    counts["操纵者"] = pd.Categorical(counts["操纵者"], categories=stack_order, ordered=True)
    counts = counts.sort_values(["所属单位", "操纵者"])
    
    fig = px.bar(
        counts,
        x="所属单位",
        y="人数",
        color="操纵者",
        text="人数",
        title="参加测试人数（按单位与技术等级）",
        color_discrete_sequence=[FIRST_OFFICER_COLOR, CAPTAIN_COLOR],
        category_orders={"操纵者": stack_order},
    )
    for trace in fig.data:
        trace.legendrank = legend_order.index(trace.name) + 1 if trace.name in legend_order else 99
    fig.update_traces(textposition="inside")
    fig.update_layout(height=430, barmode="stack", yaxis_title="人数")
    return fig


def fig_company_overall_scores(pilot_df):
    stats = pilot_df.groupby("所属单位", dropna=False)["最终得分"].mean().round(2).reset_index()
    avg_score = pilot_df["最终得分"].mean()
    fig = px.bar(
        stats,
        x="所属单位",
        y="最终得分",
        text=stats["最终得分"].map(lambda x: f"{x:.1f}"),
        title="各航司整体平均得分",
        color_discrete_sequence=[OVERALL_SCORE_COLOR],
    )
    fig.add_hline(
        y=avg_score,
        line_dash="dash",
        line_color=AVERAGE_LINE_COLOR,
    )
    fig.add_annotation(
        x=1,
        xref="paper",
        y=avg_score,
        yref="y",
        text=f"平均分: {avg_score:.2f}",
        showarrow=False,
        xanchor="right",
        yanchor="bottom",
        yshift=8,
        bgcolor="rgba(255,255,255,0.92)",
        bordercolor=AVERAGE_LINE_COLOR,
        borderwidth=1,
        font=dict(color="#333333", size=12),
    )
    fig.update_traces(textposition="outside")
    y_max = min(105, max(stats["最终得分"].max(), avg_score) + 8)
    fig.update_layout(height=430, yaxis_title="平均得分", yaxis=dict(range=[0, y_max]))
    return fig


def fig_company_role_scores(pilot_df):
    stats = (
        pilot_df.groupby(["所属单位", "操纵者"], dropna=False)["最终得分"]
        .mean()
        .round(2)
        .reset_index()
    )
    # 机长在左（绿色），副驾驶在右（蓝色）
    role_order = ["机长", "副驾驶"]
    stats["操纵者"] = pd.Categorical(stats["操纵者"], categories=role_order, ordered=True)
    stats = stats.sort_values(["所属单位", "操纵者"])
    
    fig = px.bar(
        stats,
        x="所属单位",
        y="最终得分",
        color="操纵者",
        text=stats["最终得分"].map(lambda x: f"{x:.1f}"),
        barmode="group",
        title="机长 / 副驾驶平均得分",
        color_discrete_sequence=[CAPTAIN_COLOR, FIRST_OFFICER_COLOR],
        category_orders={"操纵者": role_order},
    )
    fig.update_traces(textposition="outside")
    fig.update_layout(height=430, yaxis_title="平均得分")
    return fig


def fig_company_subject_loss(deductions, pilot_df):
    """
    各航司各科目平均失分统计
    计算逻辑：某航司某科目全部扣分值的总和 ÷ 该公司参与测试的人员数量
    科目按科目一到科目五从上到下排列
    """
    if deductions.empty or pilot_df.empty:
        return None
    
    # 获取每个航司的数据文件受测人数（按飞行员去重）
    company_people_count = company_test_counts(pilot_df)
    
    # 计算每个航司每个科目的总失分
    subject_loss_by_company = (
        deductions[deductions["科目编号"].isin(FLIGHT_SUBJECTS)]
        .groupby(["所属单位", "科目编号", "科目名称"])
        .agg(总失分=("失分", "sum"))
        .reset_index()
    )
    
    # 合并计算人均失分
    stats = subject_loss_by_company.merge(company_people_count, on="所属单位", how="left")
    stats["人均失分"] = (stats["总失分"] / stats["测试人数"].replace(0, np.nan)).fillna(0).round(2)
    
    # 添加科目显示名称（科目编号_科目名称）
    stats["科目显示"] = stats["科目编号"] + "_" + stats["科目名称"]
    
    # 按科目编号排序（科目一到科目五）
    stats["科目排序"] = stats["科目编号"].map(SUBJECT_SORT_MAP)
    stats = stats.sort_values(["所属单位", "科目排序"])
    company_order = stats["所属单位"].drop_duplicates().tolist()
    
    # 图例顺序：科目一到科目五；绘图顺序反转以保证水平分组从上到下显示为科目一到科目五。
    subject_order = [f"{no}_{name}" for no, name in SUBJECT_DEFS if no in FLIGHT_SUBJECTS]
    plot_subject_order = list(reversed(subject_order))
    subject_legend_rank = {name: idx + 1 for idx, name in enumerate(subject_order)}
    subject_color_map = {f"{no}_{name}": SUBJECT_COLORS.get(no, "#828282") for no, name in SUBJECT_DEFS if no in FLIGHT_SUBJECTS}
    
    # 创建水平条形图
    fig = px.bar(
        stats,
        y="所属单位",
        x="人均失分",
        color="科目显示",
        orientation="h",
        barmode="group",
        text=stats["人均失分"].map(lambda x: f"-{x:.1f}" if x != 0 else ""),
        title="各航司五个科目平均失分",
        category_orders={"科目显示": plot_subject_order, "所属单位": company_order},
        color_discrete_sequence=[subject_color_map.get(name, "#828282") for name in plot_subject_order],
    )
    for trace in fig.data:
        trace.legendrank = subject_legend_rank.get(trace.name, 99)
    company_avg = stats.groupby("所属单位", dropna=False)["人均失分"].mean().reindex(company_order)
    for company_idx, (company, avg_loss) in enumerate(company_avg.items()):
        fig.add_shape(
            type="line",
            x0=avg_loss,
            x1=avg_loss,
            y0=company_idx - 0.42,
            y1=company_idx + 0.42,
            xref="x",
            yref="y",
            line=dict(color=AVERAGE_LINE_COLOR, width=2, dash="dash"),
        )
        fig.add_annotation(
            x=avg_loss + 0.03,
            y=company_idx,
            text=f"-{avg_loss:.2f}",
            showarrow=False,
            xanchor="left",
            xshift=3,
            bgcolor="rgba(255,255,255,0.92)",
            bordercolor=AVERAGE_LINE_COLOR,
            borderwidth=1,
            font=dict(color="#333333", size=12),
        )
    fig.update_traces(textposition="outside")
    fig.update_layout(
        height=figure_height(stats["所属单位"].nunique(), 430, 60),
        xaxis_title="人均失分",
        yaxis_title=""
    )
    return fig


def subject_company_loss_stats(deductions, pilot_df):
    if deductions.empty or pilot_df.empty:
        return pd.DataFrame()

    company_people_count = company_test_counts(pilot_df)
    stats = (
        deductions[deductions["科目编号"].isin(FLIGHT_SUBJECTS)]
        .groupby(["所属单位", "科目编号", "科目名称"], dropna=False)
        .agg(总失分=("失分", "sum"))
        .reset_index()
        .merge(company_people_count, on="所属单位", how="left")
    )
    if stats.empty:
        return stats
    stats["人均失分"] = (stats["总失分"] / stats["测试人数"].replace(0, np.nan)).fillna(0).round(2)
    stats["科目显示"] = stats["科目编号"] + "_" + stats["科目名称"]
    stats["科目排序"] = stats["科目编号"].map(SUBJECT_SORT_MAP)
    return stats.sort_values(["科目排序", "所属单位"])


def fig_subject_company_comparison(deductions, pilot_df):
    stats = subject_company_loss_stats(deductions, pilot_df)
    if stats.empty:
        return None

    subject_order = [f"{no}_{name}" for no, name in SUBJECT_DEFS if no in FLIGHT_SUBJECTS]
    subject_avg = stats.groupby("科目显示", dropna=False)["人均失分"].mean().reindex(subject_order)
    fig = px.bar(
        stats,
        x="科目显示",
        y="人均失分",
        color="所属单位",
        barmode="group",
        text=stats["人均失分"].map(lambda x: f"-{x:.2f}" if x != 0 else ""),
        title="各科目各航司人均扣分值对比",
        category_orders={"科目显示": subject_order},
        color_discrete_sequence=DEFAULT_COLOR_SEQUENCE,
    )
    for subject_idx, (subject, avg_loss) in enumerate(subject_avg.dropna().items()):
        fig.add_shape(
            type="line",
            x0=subject_idx - 0.42,
            x1=subject_idx + 0.42,
            y0=avg_loss,
            y1=avg_loss,
            xref="x",
            yref="y",
            line=dict(color=AVERAGE_LINE_COLOR, width=2, dash="dash"),
        )
        fig.add_annotation(
            x=subject,
            y=avg_loss,
            text=f"-{avg_loss:.2f}",
            showarrow=False,
            yanchor="bottom",
            yshift=6,
            bgcolor="rgba(255,255,255,0.92)",
            bordercolor=AVERAGE_LINE_COLOR,
            borderwidth=1,
            font=dict(color="#333333", size=12),
        )
    fig.update_traces(textposition="outside", cliponaxis=False)
    y_max = max(stats["人均失分"].max(), subject_avg.max()) if not subject_avg.dropna().empty else stats["人均失分"].max()
    fig.update_layout(
        height=500,
        xaxis_title="科目",
        yaxis_title="人均扣分值",
        yaxis=dict(tickprefix="-", range=[0, y_max * 1.18 if y_max > 0 else 1]),
        margin=dict(l=60, r=60, t=80, b=80),
    )
    return fig


def build_subject_standard_company_stats(subject_deductions, pilot_df, all_items):
    if pilot_df.empty or all_items is None or all_items.empty:
        return pd.DataFrame()

    company_people_count = company_test_counts(pilot_df)
    companies = company_people_count["所属单位"].dropna().astype(str).tolist()

    base_cols = ["扣分项", "评分项目", "扣分标准", "列顺序", "标准分值"]
    join_key = "模板列" if "模板列" in all_items.columns else "扣分项"
    if join_key == "模板列":
        base_cols.append("模板列")
    base = all_items[base_cols].drop_duplicates(join_key)
    base = base.merge(pd.DataFrame({"所属单位": companies}), how="cross")

    if subject_deductions.empty:
        grouped = pd.DataFrame(columns=["所属单位", join_key, "总失分"])
    else:
        group_cols = ["所属单位", join_key] if join_key in subject_deductions.columns else ["所属单位", "扣分项"]
        grouped = (
            subject_deductions.groupby(group_cols, dropna=False)
            .agg(总失分=("失分", "sum"))
            .reset_index()
        )

    stats = (
        base.merge(grouped, on=["所属单位", join_key], how="left")
        .merge(company_people_count, on="所属单位", how="left")
    )
    stats["总失分"] = stats["总失分"].fillna(0)
    stats["人均失分"] = (stats["总失分"] / stats["测试人数"].replace(0, np.nan)).fillna(0)
    stats["显示扣分值"] = stats["人均失分"].map(lambda x: f"-{x:.2f}" if x != 0 else "")
    return stats.sort_values(["列顺序", "所属单位"])


def fig_subject_standard_company_loss(subject_deductions, pilot_df, all_items, scoring_item):
    stats = build_subject_standard_company_stats(subject_deductions, pilot_df, all_items)
    if stats.empty:
        return None

    plot_data = stats[stats["评分项目"] == scoring_item].copy()
    if plot_data.empty or plot_data["人均失分"].sum() <= 0:
        return None

    standard_order = (
        plot_data[["扣分标准", "列顺序"]]
        .drop_duplicates()
        .sort_values("列顺序")["扣分标准"]
        .tolist()
    )
    fig = px.bar(
        plot_data,
        x="扣分标准",
        y="人均失分",
        color="所属单位",
        barmode="group",
        text="显示扣分值",
        title=f"{scoring_item} 各标准航司人均扣分值",
        category_orders={"扣分标准": standard_order},
        color_discrete_sequence=DEFAULT_COLOR_SEQUENCE,
    )
    fig.update_traces(textposition="outside", cliponaxis=False)
    fig.update_layout(
        height=430,
        xaxis_title=scoring_item,
        yaxis_title="人均扣分值",
        yaxis=dict(tickprefix="-"),
        margin=dict(l=60, r=60, t=80, b=80),
    )
    return fig


def fig_subject_risk_analysis(subject_deductions, pilot_df, subject_no, subject_name, all_items):
    if pilot_df.empty or all_items is None or all_items.empty:
        return None

    base_cols = ["扣分项", "评分项目", "扣分标准", "列顺序", "标准分值"]
    join_key = "模板列" if "模板列" in all_items.columns else "扣分项"
    if join_key == "模板列":
        base_cols.append("模板列")
    base_items = all_items[base_cols].drop_duplicates(join_key).copy()
    base_items["分数权重"] = pd.to_numeric(base_items["标准分值"], errors="coerce").abs().fillna(0)
    max_weight = base_items["分数权重"].max()
    if not np.isfinite(max_weight) or max_weight <= 0:
        return None
    base_items["对应风险"] = base_items.apply(
        lambda row: risk_name_for_standard(
            row.get("科目编号", subject_no),
            row.get("评分项目", ""),
            row.get("扣分标准", ""),
            row.get("扣分项", ""),
        ),
        axis=1,
    )
    base_items["风险标签"] = (
        base_items["评分项目"].astype(str)
        + "_"
        + base_items["扣分标准"].astype(str)
        + "_"
        + base_items["对应风险"].astype(str)
    )

    company_eval_count = (
        pilot_df.groupby("所属单位", dropna=False)
        .agg(评估数据数=("评分人数", "sum"))
        .reset_index()
    )
    company_eval_count["评估数据数"] = company_eval_count["评估数据数"].replace(0, np.nan)
    companies = company_eval_count["所属单位"].dropna().astype(str).tolist()

    base = base_items.merge(pd.DataFrame({"所属单位": companies}), how="cross")
    if subject_deductions.empty:
        counts = pd.DataFrame(columns=["所属单位", join_key, "计分次数"])
    else:
        count_data = subject_deductions.copy()
        count_data["分数权重"] = pd.to_numeric(count_data["标准分值"], errors="coerce").abs().fillna(0)
        count_data["计分次数"] = np.where(
            count_data["分数权重"] > 0,
            pd.to_numeric(count_data["扣分值"], errors="coerce").abs() / count_data["分数权重"],
            1,
        )
        group_cols = ["所属单位", join_key] if join_key in count_data.columns else ["所属单位", "扣分项"]
        counts = (
            count_data.groupby(group_cols, dropna=False)
            .agg(计分次数=("计分次数", "sum"))
            .reset_index()
        )

    stats = (
        base.merge(counts, on=["所属单位", join_key], how="left")
        .merge(company_eval_count, on="所属单位", how="left")
    )
    stats["计分次数"] = stats["计分次数"].fillna(0)
    stats["风险值"] = (
        stats["分数权重"]
        * stats["计分次数"]
        * (stats["分数权重"] / (max_weight + (max_weight - stats["分数权重"])))
        / stats["评估数据数"]
    ).fillna(0)

    risk_avg = (
        stats.groupby([join_key, "扣分项", "评分项目", "扣分标准", "对应风险", "风险标签", "列顺序", "分数权重"], dropna=False)
        .agg(计分次数=("计分次数", "sum"), 评估数据数=("评估数据数", "sum"))
        .reset_index()
    )
    risk_avg["风险值"] = (
        risk_avg["分数权重"]
        * risk_avg["计分次数"]
        * (risk_avg["分数权重"] / (max_weight + (max_weight - risk_avg["分数权重"])))
        / risk_avg["评估数据数"].replace(0, np.nan)
    ).fillna(0)
    risk_avg["所属单位"] = "平权风险值"
    plot_data = pd.concat(
        [
            stats[["所属单位", "扣分项", "评分项目", "扣分标准", "对应风险", "风险标签", "列顺序", "计分次数", "评估数据数", "风险值"]],
            risk_avg[["所属单位", "扣分项", "评分项目", "扣分标准", "对应风险", "风险标签", "列顺序", "计分次数", "评估数据数", "风险值"]],
        ],
        ignore_index=True,
    )
    if plot_data["风险值"].sum() <= 0:
        return None

    label_order = base_items.sort_values("列顺序")["风险标签"].tolist()
    color_sequence = DEFAULT_COLOR_SEQUENCE + ["#1F7A3A"]
    fig = px.bar(
        plot_data,
        y="风险标签",
        x="风险值",
        color="所属单位",
        orientation="h",
        barmode="group",
        title=f"{subject_name}风险值分布一览图",
        category_orders={"风险标签": label_order},
        color_discrete_sequence=color_sequence,
        color_discrete_map={"平权风险值": "#1F7A3A"},
        custom_data=["评分项目", "扣分标准", "对应风险", "计分次数", "评估数据数"],
    )
    for trace in fig.data:
        if trace.name == "平权风险值":
            trace.marker.color = "#1F7A3A"
            trace.legendrank = 99
    fig.update_traces(
        hovertemplate=(
            "<b>%{customdata[0]}</b><br>"
            "扣分标准: %{customdata[1]}<br>"
            "对应风险: %{customdata[2]}<br>"
            "计分次数: %{customdata[3]:.2f}<br>"
            "评估数据数: %{customdata[4]:.0f}<br>"
            "风险值: %{x:.4f}<extra></extra>"
        )
    )
    fig.update_layout(
        height=figure_height(len(label_order), 520, 24, 1500),
        margin=dict(l=360, r=60, t=80, b=50),
        xaxis_title="风险值",
        yaxis_title="",
        yaxis=dict(tickfont=dict(color="#28A86B", size=11)),
    )
    return fig


def fig_loss_items_bar(data, title, top_n=None, color=None):
    """条形图：扣分值保留负数展示，但从左边对齐"""
    if data.empty:
        return None
    plot_data = data.copy()
    
    # 确保有总扣分值（保留原始负值）
    if "总扣分值" not in plot_data.columns and "总失分" in plot_data.columns:
        plot_data["总扣分值"] = negative_deduction_values(plot_data["总失分"])
    
    # 创建用于排序和显示的字段
    # 使用 abs 值进行排序（让扣分多的排在上面）
    plot_data["排序值"] = plot_data["总扣分值"].abs()
    
    if top_n:
        plot_data = plot_data.head(top_n)
        plot_data = plot_data.sort_values("排序值", ascending=True)
        category_orders = None
    elif {"科目排序", "列顺序"}.issubset(plot_data.columns):
        plot_data = plot_data.sort_values(["科目排序", "列顺序"])
        category_orders = {"扣分项": plot_data["扣分项"].tolist()}
    else:
        plot_data = plot_data.sort_values("排序值", ascending=True)
        category_orders = None
    
    # 创建条形图，使用绝对值作为 x 轴长度
    fig = px.bar(
        plot_data,
        y="扣分项",
        x="排序值",  # 使用绝对值作为条形长度（从左边开始）
        color=color,
        orientation="h",
        text=plot_data["总扣分值"].map(lambda x: f"{x:.0f}" if x != 0 else ""),  # 显示原始负值
        title=title,
        category_orders=category_orders,
        color_discrete_sequence=DEFAULT_COLOR_SEQUENCE,
    )
    
    # 自定义 hover 数据，显示原始扣分值
    fig.update_traces(
        textposition="outside", 
        cliponaxis=False,
        hovertemplate="<b>%{y}</b><br>扣分值: %{customdata:.0f}<extra></extra>",
        customdata=plot_data["总扣分值"]
    )
    
    # 更新 x 轴标签，使其显示原始负值
    x_max = plot_data["排序值"].max()
    # 创建对称的轴刻度标签（显示负值）
    tick_vals = []
    tick_labels = []
    for v in range(0, int(x_max) + 1, max(1, int(x_max // 5))):
        tick_vals.append(v)
        tick_labels.append(f"-{v}" if v > 0 else "0")
    
    fig.update_layout(
        height=figure_height(len(plot_data), 420),
        margin=dict(l=320, r=60, t=70, b=40),
        xaxis_title="总扣分值",
        yaxis_title="",
        xaxis=dict(
            tickmode="array",
            tickvals=tick_vals,
            ticktext=tick_labels,
            range=[0, x_max * 1.05]  # 从0开始，向右延伸
        )
    )
    return fig


def fig_subject_top3(deductions):
    data = deductions[deductions["科目编号"].isin(FLIGHT_SUBJECTS)].copy()
    if data.empty:
        return None, pd.DataFrame()
    grouped = (
        data.groupby(["科目编号", "科目显示", "扣分项"], dropna=False)
        .agg(总失分=("失分", "sum"), 列顺序=("列顺序", "min"), 科目排序=("科目排序", "min"))
        .reset_index()
        .sort_values(["科目排序", "总失分"], ascending=[True, False])
    )
    top3 = grouped.groupby("科目编号", group_keys=False).head(3).copy()
    top3["总扣分值"] = negative_deduction_values(top3["总失分"])
    top3["图表标签"] = top3["扣分项"]
    top3["排序值"] = top3["总扣分值"].abs()
    plot_data = top3.sort_values(["科目排序", "排序值"], ascending=[True, False])
    
    # 使用科目颜色
    subject_order = [f"{no}_{name}" for no, name in SUBJECT_DEFS if no in FLIGHT_SUBJECTS]
    color_map = {f"{no}_{name}": SUBJECT_COLORS.get(no, "#828282") for no, name in SUBJECT_DEFS if no in FLIGHT_SUBJECTS}
    
    fig = px.bar(
        plot_data,
        y="图表标签",
        x="排序值",
        color="科目显示",
        orientation="h",
        text=plot_data["总扣分值"].map(lambda x: f"{x:.0f}" if x != 0 else ""),
        title="各科目失分 TOP 3",
        category_orders={
            "图表标签": plot_data["图表标签"].tolist(),
            "科目显示": subject_order,
        },
        color_discrete_sequence=[color_map.get(s, "#828282") for s in subject_order],
    )
    fig.update_traces(
        textposition="outside", 
        cliponaxis=False,
        hovertemplate="<b>%{y}</b><br>扣分值: %{customdata:.0f}<extra></extra>",
        customdata=plot_data["总扣分值"]
    )
    x_max = plot_data["排序值"].max()
    tick_vals = []
    tick_labels = []
    for v in range(0, int(x_max) + 1, max(1, int(x_max // 5))):
        tick_vals.append(v)
        tick_labels.append(f"-{v}" if v > 0 else "0")
    fig.update_layout(
        height=figure_height(len(plot_data), 460),
        margin=dict(l=360, r=60, t=70, b=40),
        xaxis_title="总扣分值",
        yaxis_title="",
        xaxis=dict(
            tickmode="array",
            tickvals=tick_vals,
            ticktext=tick_labels,
            range=[0, x_max * 1.05]
        )
    )
    return fig, top3


def fig_subject_role_loss(subject_deductions, subject_no, all_items=None, roles=None):
    if subject_deductions.empty:
        data = pd.DataFrame(columns=["扣分项", "操纵者", "总失分", "列顺序"])
    else:
        data = (
            subject_deductions.groupby(["扣分项", "操纵者"], dropna=False)
            .agg(总失分=("失分", "sum"), 列顺序=("列顺序", "min"))
            .reset_index()
        )
    roles = roles or sorted(subject_deductions["操纵者"].dropna().astype(str).unique().tolist())
    # 确保机长在左，副驾驶在右
    role_order = ["机长", "副驾驶"]
    roles = [r for r in role_order if r in roles] + [r for r in roles if r not in role_order]
    
    if all_items is not None and not all_items.empty:
        roles = roles or ["机长", "副驾驶"]
        base = all_items[["扣分项", "列顺序"]].drop_duplicates("扣分项")
        base = base.merge(pd.DataFrame({"操纵者": roles}), how="cross")
        data = base.merge(
            data.drop(columns=[col for col in ["列顺序"] if col in data.columns]),
            on=["扣分项", "操纵者"],
            how="left",
        )
        data["总失分"] = data["总失分"].fillna(0)
    if "总扣分值" not in data.columns and "总失分" in data.columns:
        data["总扣分值"] = negative_deduction_values(data["总失分"])
    if data.empty:
        return None
    
    data["排序值"] = data["总扣分值"].abs()
    item_order = data.groupby("扣分项")["列顺序"].min().sort_values().index.tolist()
    
    # 确保操纵者顺序：机长在前，副驾驶在后
    data["操纵者"] = pd.Categorical(data["操纵者"], categories=role_order, ordered=True)
    data = data.sort_values(["扣分项", "操纵者"])
    
    fig = px.bar(
        data,
        y="扣分项",
        x="排序值",
        color="操纵者",
        orientation="h",
        category_orders={"扣分项": item_order, "操纵者": role_order},
        title=f"{subject_no} 按操纵者划分失分",
        color_discrete_sequence=[CAPTAIN_COLOR, FIRST_OFFICER_COLOR],
        text=data["总扣分值"].map(lambda x: f"{x:.0f}" if x != 0 else ""),
        custom_data=["总扣分值"],
    )
    fig.update_traces(
        textposition="outside",
        textangle=0,
        constraintext="none",
        cliponaxis=False,
        hovertemplate="<b>%{y}</b><br>扣分值: %{customdata[0]:.0f}<extra></extra>",
    )
    x_max = data["排序值"].max()
    if not np.isfinite(x_max) or x_max <= 0:
        x_max = 1
    tick_vals = []
    tick_labels = []
    for v in range(0, int(x_max) + 1, max(1, int(x_max // 5))):
        tick_vals.append(v)
        tick_labels.append(f"-{v}" if v > 0 else "0")
    fig.update_layout(
        height=figure_height(len(item_order), 430),
        margin=dict(l=320, r=100, t=70, b=40),
        xaxis_title="总扣分值",
        yaxis_title="",
        barmode="stack",
        uniformtext=dict(mode="show", minsize=10),
        xaxis=dict(
            tickmode="array",
            tickvals=tick_vals,
            ticktext=tick_labels,
            range=[0, x_max * 1.25]
        )
    )
    return fig


def fig_comprehensive_score_pie(deductions):
    data = deductions[deductions["科目编号"] == "综合考评"].copy()
    if data.empty:
        return None
    pie_data = aggregate_loss_by_item(data, ["扣分项"], sort_by_loss=False)
    pie_data = pie_data[pie_data["总失分"] > 0]
    fig = px.pie(
        pie_data,
        names="扣分项",
        values="总失分",
        title="综合考评扣分项失分占比",
        hole=0.35,
        color_discrete_sequence=DEFAULT_COLOR_SEQUENCE,
    )
    fig.update_layout(height=430)
    return fig


def build_report_figures(pilot_df, deductions, subject_melt):
    figures = []
    if not pilot_df.empty:
        figures.extend(
            [
                ("参加测试人数（按单位与技术等级）", fig_participants_by_company(pilot_df)),
                ("各航司整体平均得分", fig_company_overall_scores(pilot_df)),
                ("机长 / 副驾驶平均得分", fig_company_role_scores(pilot_df)),
                ("测试人员平均得分分布", fig_score_distribution(pilot_df)),
            ]
        )
    if not deductions.empty and not pilot_df.empty:
        figures.append(("各航司五个科目平均失分", fig_company_subject_loss(deductions, pilot_df)))
    if not deductions.empty:
        top_items = identify_weak_areas(deductions, ["扣分项"], denominator=max(len(deductions), 1))
        figures.append(("总扣分值排名前 5 的扣分项", fig_loss_items_bar(top_items, "总扣分值排名前 5 的扣分项", top_n=5)))
        fig, _ = fig_subject_top3(deductions)
        figures.append(("各科目失分 TOP 3", fig))
        figures.append(("综合考评扣分项失分占比", fig_comprehensive_score_pie(deductions)))
    return [(title, fig) for title, fig in figures if fig is not None]


def dataframe_to_doc_table(doc, df, max_rows=20):
    if df.empty:
        doc.add_paragraph("暂无数据。")
        return
    show_df = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show_df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(show_df.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in show_df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)


def add_plotly_figure_to_doc(doc, fig):
    try:
        image_bytes = fig.to_image(format="png", width=1200, height=700, scale=2)
    except Exception:
        doc.add_paragraph("当前环境未安装 Plotly 静态图片导出组件，已在 HTML 报告中保留交互式图表。")
        return False
    doc.add_picture(BytesIO(image_bytes), width=Inches(6.4))
    return True


def build_html_report(pilot_df, company_df, weak_areas, deductions, subject_melt):
    figures = build_report_figures(pilot_df, deductions, subject_melt)
    parts = [
        "<html><head><meta charset='utf-8'><title>双盲测试数据分析报告</title>",
        "<style>body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI','Microsoft YaHei',sans-serif;margin:36px;color:#222}h1,h2{margin-top:28px}table{border-collapse:collapse;width:100%;margin:12px 0 24px}th,td{border:1px solid #ddd;padding:6px 8px;font-size:13px}th{background:#f5f7fb}.metric{display:inline-block;margin-right:24px;font-size:16px}.metric b{font-size:24px}</style>",
        "</head><body>",
        "<h1>双盲测试数据分析报告</h1>",
        f"<p>生成时间：{pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}</p>",
        "<h2>一、整体情况</h2>",
        f"<p class='metric'>参加测试人数<br><b>{len(pilot_df)}</b></p>",
        f"<p class='metric'>整体平均得分<br><b>{pilot_df['最终得分'].mean():.2f}</b></p>",
        f"<p class='metric'>最高分<br><b>{pilot_df['最终得分'].max():.1f}</b></p>",
        f"<p class='metric'>最低分<br><b>{pilot_df['最终得分'].min():.1f}</b></p>",
        "<h2>二、各单位统计</h2>",
        company_df.to_html(index=False, escape=False) if not company_df.empty else "<p>暂无数据。</p>",
        "<h2>三、核心图表</h2>",
    ]
    include_plotlyjs = True
    for title, fig in figures:
        parts.append(f"<h3>{title}</h3>")
        parts.append(pio.to_html(fig, full_html=False, include_plotlyjs=include_plotlyjs))
        include_plotlyjs = False
    parts.extend(
        [
            "<h2>四、薄弱环节 TOP 10</h2>",
            display_loss_table(weak_areas.head(10)).to_html(index=False, escape=False) if not weak_areas.empty else "<p>未发现扣分项。</p>",
            "</body></html>",
        ]
    )
    return "\n".join(parts).encode("utf-8")


def build_simple_report(pilot_df, company_df, weak_areas, deductions=None, subject_melt=None):
    if not DOCX_AVAILABLE:
        return None
    deductions = deductions if deductions is not None else pd.DataFrame()
    subject_melt = subject_melt if subject_melt is not None else pd.DataFrame()

    doc = Document()
    title = doc.add_heading("双盲测试数据分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间：{pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("一、整体情况", level=1)
    doc.add_paragraph(f"参加测试人数：{len(pilot_df)} 人")
    doc.add_paragraph(f"整体平均得分：{pilot_df['最终得分'].mean():.2f} 分")
    doc.add_paragraph(
        f"最高分：{pilot_df['最终得分'].max():.1f} 分，最低分：{pilot_df['最终得分'].min():.1f} 分"
    )

    if not company_df.empty:
        doc.add_heading("二、各单位统计", level=1)
        dataframe_to_doc_table(doc, company_df)

    doc.add_heading("三、核心图表", level=1)
    for fig_title, fig in build_report_figures(pilot_df, deductions, subject_melt):
        doc.add_heading(fig_title, level=2)
        add_plotly_figure_to_doc(doc, fig)

    doc.add_heading("四、薄弱环节", level=1)
    if weak_areas.empty:
        doc.add_paragraph("未发现扣分项。")
    else:
        show_cols = [col for col in ["扣分项", "扣分次数", "总扣分值", "出现率"] if col in weak_areas.columns]
        dataframe_to_doc_table(doc, weak_areas[show_cols].head(10))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


st.title("飞行员双盲测试评估分析平台（空客）")
st.caption("得分按飞行员取检查员平均，扣分项按检查员原始记录统计。")

with st.sidebar:
    st.markdown("### 数据上传")
    uploaded_files = st.file_uploader(
        "上传一个或多个Excel文件",
        type=["xlsx", "xls"],
        accept_multiple_files=True,
    )
    st.markdown("### 评分规则")
    st.markdown("- 基础分：100分\n- 飞行员得分：各检查员评分平均值\n- 扣分统计：使用检查员原始扣分记录")


if uploaded_files:
    all_data_list = []
    all_meta = OrderedDict()
    all_deduction_cols = []
    success_count = 0

    for file in uploaded_files:
        with st.spinner(f"正在处理：{file.name}"):
            df, col_meta, deduction_cols = load_template(file.getvalue(), file.name)
            if not df.empty:
                all_data_list.append(df)
                all_meta.update(col_meta)
                all_deduction_cols.extend([col for col in deduction_cols if col not in all_deduction_cols])
                success_count += 1
                st.success(f"{file.name}：读取 {len(df)} 条检查记录")
            else:
                st.warning(f"{file.name}：未能提取有效数据")

    if not all_data_list:
        st.error("没有成功加载任何文件，请检查文件格式是否符合模板要求。")
        st.stop()

    raw_data = pd.concat(all_data_list, ignore_index=True)
    raw_data, pilot_data, deductions, subject_melt, subject_to_cols = prepare_analysis(
        raw_data, all_meta, all_deduction_cols
    )

    weak_areas = identify_weak_areas(deductions, ["扣分项"], denominator=len(raw_data[~raw_data["是否平均行"]]))
    company_df = company_stats(pilot_data)

    with st.sidebar:
        companies = ["全部"] + sorted(pilot_data["所属单位"].dropna().astype(str).unique().tolist())
        selected_company = st.selectbox("单位筛选", companies)

    filtered_pilots = pilot_data.copy()
    filtered_deductions = deductions.copy()
    filtered_subject_melt = subject_melt.copy()

    if selected_company != "全部":
        filtered_pilots = filtered_pilots[filtered_pilots["所属单位"] == selected_company]
        filtered_deductions = filtered_deductions[filtered_deductions["所属单位"] == selected_company]
        filtered_subject_melt = filtered_subject_melt[filtered_subject_melt["所属单位"] == selected_company]

    st.success(
        f"成功加载 {len(pilot_data)} 位测试人员、{len(raw_data[~raw_data['是否平均行']])} 条检查员评分记录，"
        f"识别 {len(all_deduction_cols)} 个扣分列。"
    )

    tab1, tab2, tab3, tab4 = st.tabs(["整体情况分析", "科目分析", "个人详情", "报告输出"])

    with tab1:
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("参加测试人数", len(filtered_pilots))
        col2.metric("平均分", f"{filtered_pilots['最终得分'].mean():.1f}")
        col3.metric("最高分", f"{filtered_pilots['最终得分'].max():.1f}")
        col4.metric("最低分", f"{filtered_pilots['最终得分'].min():.1f}")

        st.plotly_chart(fig_participants_by_company(filtered_pilots), use_container_width=True)

        score_col1, score_col2 = st.columns(2)
        with score_col1:
            st.plotly_chart(fig_company_overall_scores(filtered_pilots), use_container_width=True)
        with score_col2:
            st.plotly_chart(fig_company_role_scores(filtered_pilots), use_container_width=True)

        # 显示主图：所有飞行员的分数分布
        st.plotly_chart(fig_score_distribution(filtered_pilots), use_container_width=True)

        # 显示机长和副驾驶的分数分布
        if not filtered_pilots.empty:
            role_col1, role_col2 = st.columns(2)
            with role_col1:
                fig_captain = fig_score_distribution_by_role(filtered_pilots, "机长")
                if fig_captain:
                    st.plotly_chart(fig_captain, use_container_width=True)
                else:
                    st.info("无机长数据")
            with role_col2:
                fig_first_officer = fig_score_distribution_by_role(filtered_pilots, "副驾驶")
                if fig_first_officer:
                    st.plotly_chart(fig_first_officer, use_container_width=True)
                else:
                    st.info("无副驾驶数据")

        st.markdown("#### 各单位得分统计")
        shown_company_df = company_stats(filtered_pilots)
        st.dataframe(shown_company_df, use_container_width=True, hide_index=True)

    with tab2:
        if filtered_deductions.empty:
            st.success("未发现扣分项。")
        else:
            st.markdown("#### 各航司平均失分统计")
            fig = fig_company_subject_loss(filtered_deductions, filtered_pilots)
            if fig:
                st.plotly_chart(fig, use_container_width=True)

            st.markdown("#### 各科目各航司人均扣分值对比")
            fig = fig_subject_company_comparison(filtered_deductions, filtered_pilots)
            if fig:
                st.plotly_chart(fig, use_container_width=True)

            st.markdown("#### 失分 TOP 5 扣分项统计")
            flight_deductions = filtered_deductions[filtered_deductions["科目编号"].isin(FLIGHT_SUBJECTS)].copy()
            top_items = identify_weak_areas(flight_deductions, ["扣分项"], denominator=len(raw_data))
            fig = fig_loss_items_bar(top_items, "总扣分值排名前 5 的扣分项", top_n=5)
            if fig:
                st.plotly_chart(fig, use_container_width=True)
            st.dataframe(display_loss_table(top_items.head(5)), use_container_width=True, hide_index=True)

            st.markdown("#### 各科目失分 TOP 3 统计")
            fig, top3 = fig_subject_top3(filtered_deductions)
            if fig:
                st.plotly_chart(fig, use_container_width=True)
                st.dataframe(display_loss_table(top3), use_container_width=True, hide_index=True)

            st.markdown("#### 各科目分析")
            for subject_no in FLIGHT_SUBJECTS:
                subject_name = dict(SUBJECT_DEFS).get(subject_no, subject_no)
                subject_deductions = filtered_deductions[filtered_deductions["科目编号"] == subject_no].copy()
                subject_all_items = all_deduction_items(all_meta, subject_no)
                with st.expander(f"{subject_no} {subject_name}", expanded=(subject_no == "科目一")):
                    if subject_deductions.empty and subject_all_items.empty:
                        st.info("该科目暂无扣分记录。")
                        continue

                    inner_tab1, inner_tab2, inner_tab3, inner_tab4 = st.tabs(["整体分析", "按操纵者划分", "各航司失分情况", "风险值分析"])

                    with inner_tab1:
                        item_loss = aggregate_loss_by_item(
                            subject_deductions,
                            ["扣分项"],
                            sort_by_loss=False,
                            all_items=subject_all_items,
                        )
                        fig = fig_loss_items_bar(item_loss, f"{subject_no} 全部评分项目及扣分项总扣分值")
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)

                        if not subject_all_items.empty and "评分项目" in subject_all_items.columns:
                            st.markdown("##### 各评分项目下扣分标准航司人均扣分值")
                            scoring_items = (
                                subject_all_items[["评分项目", "列顺序"]]
                                .drop_duplicates("评分项目")
                                .sort_values("列顺序")["评分项目"]
                                .tolist()
                            )
                            shown_standard_chart = False
                            for scoring_item in scoring_items:
                                fig = fig_subject_standard_company_loss(
                                    subject_deductions,
                                    filtered_pilots,
                                    subject_all_items,
                                    scoring_item,
                                )
                                if fig:
                                    st.plotly_chart(fig, use_container_width=True)
                                    shown_standard_chart = True
                            if not shown_standard_chart:
                                st.info("该科目暂无可绘制的标准扣分数据。")
                        st.dataframe(display_loss_table(item_loss), use_container_width=True, hide_index=True)

                    with inner_tab2:
                        roles = sorted(filtered_pilots["技术等级"].dropna().astype(str).unique().tolist())
                        fig = fig_subject_role_loss(subject_deductions, subject_no, subject_all_items, roles)
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)

                    with inner_tab3:
                        for company in sorted(filtered_pilots["所属单位"].dropna().astype(str).unique()):
                            company_data = subject_deductions[subject_deductions["所属单位"].astype(str) == company]
                            company_loss = aggregate_loss_by_item(
                                company_data,
                                ["扣分项"],
                                sort_by_loss=False,
                                all_items=subject_all_items,
                            )
                            fig = fig_loss_items_bar(company_loss, f"{company} - {subject_no} 总扣分值")
                            if fig:
                                st.plotly_chart(fig, use_container_width=True)

                    with inner_tab4:
                        fig = fig_subject_risk_analysis(
                            subject_deductions,
                            filtered_pilots,
                            subject_no,
                            subject_name,
                            subject_all_items,
                        )
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)
                        else:
                            st.info("该科目暂无可绘制的风险值数据。")

            st.markdown("#### 综合考评得分分析")
            fig = fig_comprehensive_score_pie(filtered_deductions)
            if fig:
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("暂无综合考评数据。")

    with tab3:
        st.markdown("#### 个人 / 单位 / 技术等级筛选")
        detail_col1, detail_col2 = st.columns(2)
        with detail_col1:
            detail_companies = ["全部"] + sorted(filtered_pilots["所属单位"].dropna().astype(str).unique().tolist())
            detail_company = st.selectbox("个人详情单位筛选", detail_companies, key="detail_company")
        with detail_col2:
            detail_roles = ["全部"] + sorted(filtered_pilots["技术等级"].dropna().astype(str).unique().tolist())
            detail_role = st.selectbox("技术等级筛选", detail_roles, key="detail_role")

        detail_pilots = filtered_pilots.copy()
        detail_deductions = filtered_deductions.copy()
        if detail_company != "全部":
            detail_pilots = detail_pilots[detail_pilots["所属单位"] == detail_company]
            detail_deductions = detail_deductions[detail_deductions["所属单位"] == detail_company]
        if detail_role != "全部":
            detail_pilots = detail_pilots[detail_pilots["技术等级"] == detail_role]
            detail_deductions = detail_deductions[detail_deductions["技术等级"] == detail_role]

        metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)
        metric_col1.metric("筛选人数", len(detail_pilots))
        metric_col2.metric("平均分", f"{detail_pilots['最终得分'].mean():.1f}" if not detail_pilots.empty else "-")
        metric_col3.metric("最高分", f"{detail_pilots['最终得分'].max():.1f}" if not detail_pilots.empty else "-")
        metric_col4.metric("最低分", f"{detail_pilots['最终得分'].min():.1f}" if not detail_pilots.empty else "-")

        if not detail_pilots.empty:
            chart_col1, chart_col2 = st.columns(2)
            with chart_col1:
                st.plotly_chart(fig_score_distribution(detail_pilots, "筛选范围得分分布"), use_container_width=True)
            with chart_col2:
                detail_top = identify_weak_areas(detail_deductions, ["扣分项"], denominator=max(len(detail_deductions), 1))
                fig = fig_loss_items_bar(detail_top, "筛选范围扣分项 TOP 5", top_n=5)
                if fig:
                    st.plotly_chart(fig, use_container_width=True)

        display_cols = [
            "姓名",
            "所属单位",
            "技术等级",
            "操纵者",
            "评分人数",
            "最终得分",
            "平均失分",
            "平均扣分项数量",
        ]
        display_cols = [col for col in display_cols if col in detail_pilots.columns]
        st.dataframe(
            detail_pilots[display_cols].sort_values("最终得分", ascending=False),
            use_container_width=True,
            hide_index=True,
        )

        if not detail_deductions.empty:
            st.markdown("#### 个人扣分明细查询")
            pilot_list = detail_pilots["姓名"].dropna().astype(str).unique().tolist()
            if pilot_list:
                selected_pilot = st.selectbox("选择飞行员", pilot_list)
                pilot_deductions = detail_deductions[detail_deductions["姓名"] == selected_pilot]
                if not pilot_deductions.empty:
                    st.dataframe(
                        pilot_deductions[
                            ["检查员", "科目编号", "科目名称", "评分项目", "扣分标准", "扣分值", "扣分项"]
                        ],
                        use_container_width=True,
                        hide_index=True,
                    )
                else:
                    st.success("该飞行员无扣分项。")

    with tab4:
        st.markdown("### 下载报告")
        html_report = build_html_report(pilot_data, company_df, weak_areas, deductions, subject_melt)
        st.download_button(
            "下载 HTML 图文报告",
            html_report,
            f"双盲测试图文报告_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.html",
            mime="text/html",
        )

        report_bio = build_simple_report(pilot_data, company_df, weak_areas, deductions, subject_melt)
        if report_bio is not None:
            st.download_button(
                "下载 Word 报告",
                report_bio,
                f"双盲测试报告_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.info("当前 Python 环境未安装 python-docx，Word 报告下载暂不可用。HTML 图文报告和 CSV 导出仍可使用。")

        st.download_button(
            "下载飞行员平均得分 CSV",
            pilot_data.to_csv(index=False, encoding="utf-8-sig"),
            f"飞行员平均得分_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
        )

        if not deductions.empty:
            st.download_button(
                "下载检查员原始扣分明细 CSV",
                deductions.to_csv(index=False, encoding="utf-8-sig"),
                f"检查员原始扣分明细_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
            )

        st.markdown("---")
        st.markdown("### 报告覆盖维度")
        st.markdown(
            """
            - 整体情况：参加测试人数、平均得分、机长 / 副驾驶得分对比
            - 科目分析：航司平均失分、TOP 扣分项、各科目拆分分析、综合考评
            - 个人详情：按飞行员平均得分排名，查看检查员原始扣分明细
            - 导出功能：平均得分 CSV、原始扣分明细 CSV、Word 报告
            """
        )

else:
    st.info("请先从左侧上传双盲测试 Excel 数据。")
    with st.expander("文件格式说明"):
        st.markdown(
            """
            支持按《华东飞行员双盲测试数据采集表》模板填写的 Excel 文件。系统会自动识别模板中的科目、评分项目和扣分标准，并将扣分项命名为：

            `科目X_评分项目_扣分项`

            示例：`科目一_高度偏差（偏差持续5秒，加倍）_±60-80ft（含）`
            """
        )

st.markdown("---")
